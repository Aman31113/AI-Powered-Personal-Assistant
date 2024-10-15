import pyttsx3
import speech_recognition 
import os
import pyautogui
import pygetwindow as gw
import datetime
from docx import Document

engine = pyttsx3.init("sapi5")
voices = engine.getProperty("voices")
engine.setProperty("voice", voices[0].id)
rate = engine.setProperty("rate",170)
volume = engine.setProperty("volume",1000)

def speak(audio):
    engine.say(audio)
    engine.runAndWait()

def takeCommand():
    r = speech_recognition.Recognizer()
    with speech_recognition.Microphone() as source:
        print("Listening.....")
        r.pause_threshold = 1
        r.energy_threshold = 300
        audio = r.listen(source,0,4)
        r.adjust_for_ambient_noise(source) 
        

    try:
        print("Understanding..")
        query  = r.recognize_google(audio,language='en-in')
        print(f"You Said: {query}\n")
    except Exception as e:
        print("Say that again")
        return "None"
    return query


def record_and_write_to_word():
    # Initialize the recognizer
    recognizer = speech_recognition.Recognizer()

    # Use the default microphone as the audio source
    with speech_recognition.Microphone() as source:
        print("Listening...")
        speak("start recording sir")

        # Adjust the energy threshold for ambient noise levels
        recognizer.adjust_for_ambient_noise(source, duration=1)

        # Record audio from the microphone
        audio = recognizer.listen(source)

        print("Recognizing...")

        try:
            # Use the Google Web Speech API to recognize the audio
            text = recognizer.recognize_google(audio)

            # Get the current date and time
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")


            # Create a new Document object
            doc = Document()

            # Add a heading with the current date and time
            doc.add_heading(f"Transcription - {timestamp}", level=1)

            # Add the recognized text to the document
            doc.add_paragraph(text)

            # Save the document to a Word file
            doc.save("transcription.docx")
            print("Transcription saved to transcription.docx")

        except speech_recognition.UnknownValueError:
            print("Google Web Speech API could not understand audio")
        except speech_recognition.RequestError as e:
            print(f"Could not request results from Google Web Speech API; {e}")
            
        
def record_and_write_to_notepad():
    # Initialize the recognizer
    recognizer = speech_recognition.Recognizer()

    # Use the default microphone as the audio source
    with speech_recognition.Microphone() as source:
        print("Listening...")
        speak("recording start sir")

        # Adjust the energy threshold for ambient noise levels
        recognizer.adjust_for_ambient_noise(source, duration=1)

        # Record audio from the microphone
        audio = recognizer.listen(source)

        print("Recognizing...")

        try:
            # Use the Google Web Speech API to recognize the audio
            text = recognizer.recognize_google(audio)
            timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S\n")


            # Write the recognized text to a Notepad file
            with open("speech_to_text.txt", "a") as file:
                file.write(f"{timestamp}: {text}\n\n\n")
                print("Transcription saved to speech_to_text.txt")
                speak("file saved , sir")

        except speech_recognition.UnknownValueError:
            print("Google Web Speech API could not understand audio")
        except speech_recognition.RequestError as e:
            print(f"Could not request results from Google Web Speech API; {e}")



# if __name__ == "__main__":
#     record_and_write_to_word()

